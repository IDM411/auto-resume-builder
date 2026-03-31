from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, Field
from typing import List
import json
import pdfplumber
import docx
import io
import requests 
from fpdf import FPDF



# Initialize the server
app = FastAPI()

# --- 1. THE AI SCHEMAS ---
class Education(BaseModel):
    school: str = Field(description="Name of the university or college")
    degree: str = Field(description="Degree and Major")
    location: str = Field(description="City, State")
    date: str = Field(description="Graduation Date")

class Project(BaseModel):
    name: str = Field(description="Name of the project")
    technologies: str = Field(description="Tools used")
    date: str = Field(description="Date range")
    bullets: List[str]

class Experience(BaseModel):
    company: str
    title: str
    location: str
    date: str
    bullets: List[str]
    explanation: str = Field(description="Why this was changed.")

class TailoredResume(BaseModel):
    full_name: str
    contact_info: str
    summary: str
    education: List[Education]
    experience: List[Experience]
    projects: List[Project]
    skills: List[str]

class ResumeAnalysis(BaseModel):
    original_ats_score: int = Field(description="Score from 0-100 of the original resume vs JD")
    tailored_ats_score: int = Field(default=100, description="The new score after tailoring")
    critique: str = Field(description="Overall feedback on the candidate's fit")
    optimization_report: List[str] = Field(description="List of 3-5 specific ways the JD was used to enhance the resume")
    tailored_resume: TailoredResume


# --- 2. FILE READER ---
def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    text = ""
    try:
        if filename.lower().endswith('.pdf'):
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted: text += extracted + "\n"
        elif filename.lower().endswith('.docx'):
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs: text += para.text + "\n"
        return text
    except Exception as e:
        raise ValueError(f"Could not read the file: {str(e)}")


# --- 3. CORE ROUTES ---
@app.get("/")
async def serve_homepage():
    return FileResponse("index.html")

@app.post("/api/analyze-and-tailor")
async def analyze_and_tailor(resume_file: UploadFile = File(...), job_description: str = Form(...)):
    try:
        file_bytes = await resume_file.read()
        master_resume_text = extract_text_from_file(file_bytes, resume_file.filename)

        schema_blueprint = json.dumps(ResumeAnalysis.model_json_schema(), indent=2)

        system_instruction = f"""
### ROLE
You are an Elite Technical Recruiter. Your goal is to re-architect this resume for a 100/100 ATS match.

### CATEGORIZATION LOGIC (CRITICAL)
- **NO DUPLICATION:** Every item from the Master Resume must appear in **exactly one** section. 
- **EXPERIENCE:** Use this for paid work, internships, or long-term roles with a specific company title.
- **PROJECTS:** Use this for student hackathons (like KnightHacks), personal labs, or independent game dev (like Mindbreak) if they aren't part of a formal job.
- **DECISION RULE:** If an item could fit in both, prioritize **Experience** only if it was a paid/formal role. Otherwise, put it in **Projects**. DO NOT REPEAT THE SAME CONTENT.

### OBJECTIVES
1. Analyze JD vs Master Resume.
2. Calculate scores and optimization strategy.
3. Rewrite bullets using the STAR method.
4. Output valid JSON only.
--- SCHEMA ---
{schema_blueprint}
"""
        full_prompt = f"{system_instruction}\n\nMASTER RESUME:\n{master_resume_text}\n\nJOB DESCRIPTION:\n{job_description}"

        # TWEAK: Using Qwen 32B with Keep-Alive to stay in your 64GB RAM
        payload = {
            "model": "qwen2.5:32b", 
            "prompt": full_prompt,
            "format": "json",       
            "stream": False,
            "options": { "temperature": 0.1, "num_ctx": 12288 },
            "keep_alive": -1  # <--- Keep model loaded in memory
        }

        # TWEAK: Increased timeout to 900s (15 mins) for your workstation
        response = requests.post("http://localhost:11434/api/generate", json=payload, timeout=900)
        
        response_data = response.json()
        raw_ai_text = response_data.get("response", "{}").strip()

        # TWEAK: Markdown cleaner (Removes ```json ... ``` if the AI adds it)
        if raw_ai_text.startswith("```"):
            raw_ai_text = raw_ai_text.split("\n", 1)[1].rsplit("\n", 1)[0].strip()
            if raw_ai_text.startswith("json"):
                raw_ai_text = raw_ai_text[4:].strip()

        print(f"\n--- AI OUTPUT ---\n{raw_ai_text}\n-----------------\n")

        return {
            "original_text": master_resume_text,
            "analysis": json.loads(raw_ai_text)
        }

    except Exception as e:
        print(f"\n--- CRITICAL AI ERROR ---\n{str(e)}\n-------------------------\n")
        raise HTTPException(status_code=500, detail=str(e))


# --- 4. DOWNLOAD ROUTES ---
def clean_text(text):
    if not text: return ""
    text = str(text).replace('“', '"').replace('”', '"').replace('‘', "'").replace('’', "'")
    text = text.replace('–', '-').replace('—', '-').replace('•', '-')
    return text.encode('ascii', 'ignore').decode('ascii')

@app.post("/api/download-docx")
async def generate_word_doc(resume_data: TailoredResume):
    try:
        doc = docx.Document()
        doc.add_heading(clean_text(resume_data.full_name), level=0)
        doc.add_paragraph(clean_text(resume_data.contact_info))
        
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(clean_text(resume_data.summary))
        
        if resume_data.education:
            doc.add_heading('Education', level=1)
            for edu in resume_data.education:
                p = doc.add_paragraph()
                p.add_run(f"{clean_text(edu.school)}").bold = True
                p.add_run(f" | {clean_text(edu.degree)} ({clean_text(edu.date)})").italic = True

        doc.add_heading('Experience', level=1)
        for job in resume_data.experience:
            p = doc.add_paragraph()
            p.add_run(f"{clean_text(job.title)}").bold = True
            p.add_run(f" | {clean_text(job.company)}").italic = True
            for bullet in job.bullets:
                doc.add_paragraph(clean_text(bullet), style='List Bullet')

        if resume_data.projects:
            doc.add_heading('Projects', level=1)
            for proj in resume_data.projects:
                p = doc.add_paragraph()
                p.add_run(f"{clean_text(proj.name)}").bold = True
                p.add_run(f" | {clean_text(proj.technologies)}").italic = True
                for bullet in proj.bullets:
                    doc.add_paragraph(clean_text(bullet), style='List Bullet')
                    
        doc.add_heading('Skills', level=1)
        doc.add_paragraph(clean_text(", ".join(resume_data.skills)))
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return StreamingResponse(
            file_stream, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Tailored_Resume.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/download-pdf")
async def generate_pdf(resume_data: TailoredResume):
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        pdf.set_font("helvetica", "B", 16)
        pdf.cell(0, 10, clean_text(resume_data.full_name), ln=True, align="C")
        pdf.set_font("helvetica", "", 10)
        pdf.cell(0, 5, clean_text(resume_data.contact_info), ln=True, align="C")
        pdf.ln(5)

        pdf.set_font("helvetica", "B", 12)
        pdf.cell(0, 8, "Summary", ln=True)
        pdf.set_font("helvetica", "", 11)
        pdf.write(6, clean_text(resume_data.summary) + "\n\n")

        if resume_data.education:
            pdf.set_font("helvetica", "B", 12)
            pdf.cell(0, 8, "Education", ln=True)
            for edu in resume_data.education:
                pdf.set_font("helvetica", "B", 11)
                pdf.write(6, clean_text(edu.school) + " ")
                pdf.set_font("helvetica", "", 11)
                pdf.write(6, clean_text(f"| {edu.degree} | {edu.date}\n"))
            pdf.write(6, "\n")

        pdf.set_font("helvetica", "B", 12)
        pdf.cell(0, 8, "Experience", ln=True)
        for job in resume_data.experience:
            pdf.set_font("helvetica", "B", 11)
            pdf.write(7, clean_text(f"{job.title} | {job.company}") + "\n")
            pdf.set_font("helvetica", "", 11)
            for bullet in job.bullets:
                pdf.write(6, f"- {clean_text(bullet)}\n")
            pdf.write(6, "\n")

        if resume_data.projects:
            pdf.set_font("helvetica", "B", 12)
            pdf.cell(0, 8, "Projects", ln=True)
            for proj in resume_data.projects:
                pdf.set_font("helvetica", "B", 11)
                pdf.write(7, clean_text(f"{proj.name} | {proj.technologies}") + "\n")
                pdf.set_font("helvetica", "", 11)
                for bullet in proj.bullets:
                    pdf.write(6, f"- {clean_text(bullet)}\n")
                pdf.write(6, "\n")

        pdf.set_font("helvetica", "B", 12)
        pdf.cell(0, 8, "Skills", ln=True)
        pdf.set_font("helvetica", "", 11)
        pdf.write(6, clean_text(", ".join(resume_data.skills)) + "\n")

        pdf_bytes = bytes(pdf.output())
        return StreamingResponse(
            io.BytesIO(pdf_bytes), 
            media_type="application/pdf",
            headers={
                "Content-Disposition": "attachment; filename=Tailored_Resume.pdf",
                "Content-Length": str(len(pdf_bytes))
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/generate-latex")
async def generate_latex(resume_data: TailoredResume):
    try:
        def tex_escape(text):
            if not text: return ""
            replacements = {'&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#', '_': r'\_', '{': r'\{', '}': r'\}'}
            for key, val in replacements.items(): text = text.replace(key, val)
            return text

        latex_template = r"""\documentclass[letterpaper,11pt]{article}
\usepackage{latexsym}
\usepackage[empty]{fullpage}
\usepackage{titlesec}
\usepackage{marvosym}
\usepackage[usenames,dvipsnames]{color}
\usepackage{verbatim}
\usepackage{enumitem}
\usepackage[hidelinks]{hyperref}
\usepackage{fancyhdr}
\usepackage[english]{babel}
\usepackage{tabularx}
\input{glyphtounicode}

\pagestyle{fancy}
\fancyhf{} 
\renewcommand{\headrulewidth}{0pt}
\renewcommand{\footrulewidth}{0pt}

\addtolength{\oddsidemargin}{-0.5in}
\addtolength{\evensidemargin}{-0.5in}
\addtolength{\textwidth}{1in}
\addtolength{\topmargin}{-0.3in}
\addtolength{\textheight}{0.8in}
\urlstyle{same}
\raggedbottom
\raggedright
\setlength{\tabcolsep}{0in}

\titleformat{\section}{\vspace{-2pt}\scshape\raggedright\large}{}{0em}{}[\color{black}\titlerule \vspace{-2pt}]
\pdfgentounicode=1

\newcommand{\resumeItem}[1]{\item\small{{#1 \vspace{1pt}}}}
\newcommand{\resumeSubheading}[4]{
  \vspace{2pt}\item
    \begin{tabular*}{0.97\textwidth}[t]{l@{\extracolsep{\fill}}r}
      \textbf{#1} & #2 \\
      \textit{\small#3} & \textit{\small #4} \\
    \end{tabular*}\vspace{-2pt}
}
\newcommand{\resumeProjectHeading}[2]{
    \item
    \begin{tabular*}{0.97\textwidth}{l@{\extracolsep{\fill}}r}
      \small#1 & #2 \\
    \end{tabular*}\vspace{-2pt}
}
\newcommand{\resumeSubHeadingListStart}{\begin{itemize}[leftmargin=0.15in, label={}]}
\newcommand{\resumeSubHeadingListEnd}{\end{itemize}}
\newcommand{\resumeItemListStart}{\begin{itemize}}
\newcommand{\resumeItemListEnd}{\end{itemize}\vspace{2pt}}

\begin{document}
\begin{center}
    \textbf{\Huge \scshape [[FULL_NAME]]} \\ \vspace{4pt}
    \small [[CONTACT_INFO]]
\end{center}

\section{Education}
  \resumeSubHeadingListStart
[[EDUCATION_BLOCKS]]
  \resumeSubHeadingListEnd

[[EXPERIENCE_SECTION]]
[[PROJECT_SECTION]]

\section{Technical Skills}
 \begin{itemize}[leftmargin=0.15in, label={}]
    \small{\item{
     \textbf{Core Competencies}{: [[SKILLS]]}
    }}
 \end{itemize}
\end{document}
"""
        edu_blocks = ""
        for edu in resume_data.education:
            edu_blocks += f"    \\resumeSubheading\n      {{{tex_escape(edu.school)}}}{{{tex_escape(edu.date)}}}\n      {{{tex_escape(edu.degree)}}}{{{tex_escape(edu.location)}}}\n"

        exp_blocks = ""
        exp_section = ""
        if resume_data.experience:
            for job in resume_data.experience:
                exp_blocks += f"    \\resumeSubheading\n      {{{tex_escape(job.title)}}}{{{tex_escape(job.date)}}}\n      {{{tex_escape(job.company)}}}{{{tex_escape(job.location)}}}\n      \\resumeItemListStart\n"
                for bullet in job.bullets: exp_blocks += f"        \\resumeItem{{{tex_escape(bullet)}}}\n"
                exp_blocks += f"      \\resumeItemListEnd\n"
            exp_section = "\\section{Experience}\n  \\resumeSubHeadingListStart\n" + exp_blocks + "  \\resumeSubHeadingListEnd\n"

        proj_blocks = ""
        proj_section = ""
        if resume_data.projects:
            for proj in resume_data.projects:
                proj_blocks += f"    \\resumeProjectHeading\n      {{\\textbf{{{tex_escape(proj.name)}}} $|$ \\emph{{{tex_escape(proj.technologies)}}}}}{{{tex_escape(proj.date)}}}\n      \\resumeItemListStart\n"
                for bullet in proj.bullets: proj_blocks += f"        \\resumeItem{{{tex_escape(bullet)}}}\n"
                proj_blocks += f"      \\resumeItemListEnd\n"
            proj_section = "\\section{Projects}\n  \\resumeSubHeadingListStart\n" + proj_blocks + "  \\resumeSubHeadingListEnd\n"

        final_tex = latex_template.replace("[[FULL_NAME]]", tex_escape(resume_data.full_name))
        final_tex = final_tex.replace("[[CONTACT_INFO]]", tex_escape(resume_data.contact_info))
        final_tex = final_tex.replace("[[EDUCATION_BLOCKS]]", edu_blocks)
        final_tex = final_tex.replace("[[EXPERIENCE_SECTION]]", exp_section)
        final_tex = final_tex.replace("[[PROJECT_SECTION]]", proj_section)
        final_tex = final_tex.replace("[[SKILLS]]", tex_escape(", ".join(resume_data.skills)))

        return {"latex_code": final_tex}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))